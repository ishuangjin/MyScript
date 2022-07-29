import docx
from win32com import client
import pandas as pd


def doc_docx(doc_path):
    # 路径
    new_path = doc_path.replace('doc', 'docx')
    # 获取Word应用程序对象
    word_app = client.Dispatch('Word.Application')
    # 打开对应的Word文档
    doc = word_app.Documents.Open(doc_path)
    # 另存到docx_path中，12表示docx文件格式
    doc.SaveAs(new_path, 12)
    # 关闭Word文档
    doc.Close()
    # 退出软件
    word_app.Quit()
    return new_path


if __name__ == '__main__':
    doc_path = r'E:\pythonProject\读word中的表格写入excel\POC测试用例.docx'
    try:
        doc = docx.Document(doc_path)
    except Exception:
        doc = docx.Document(doc_docx(doc_path))

    tables = doc.tables
    excel_path = doc_path.replace('docx', 'xlsx')
    list_ = []  # 初始化一个空列表，用来装后面的dict_
    for table in tables:  # 循环所有的表格列表
        dict_ = {}
        for i in range(0, len(table.rows)):  # 后续的内容格式不固定，所以循环获取
            dict_[table.cell(i, 0).text] = table.cell(i, 1).text
        list_.append(dict_)
    df = pd.DataFrame(list_)

    # 重新设置一下行列索引的名字
    df.columns = ["用例名称", "测试说明", "前置条件", "测试步骤", "预期结果", "实际结果", "备注", "用例标识"]  # , "功能需求追踪", "用例编制人", "测试执行人", "测试结论"
    df = df.reindex(columns=['用例标识', '用例名称', '功能需求追踪', '测试说明', '前置条件', '测试步骤', '预期结果', '用例编制人', '实际结果', '测试结论', '测试执行人', '备注'])
    df.to_excel(excel_path, index=None)

