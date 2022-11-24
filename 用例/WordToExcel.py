#!/usr/bin/env python3
# encoding:utf-8
'''
@Author: ishuangjin
@WebSite: blog.ishuangjin.cn
@QQ: 1525053461
@Mail: ishuangjin@foxmail.com
@Date: 2022-11-07 12:57:38
@LastEditTime: 2022-11-24 15:34:08
@FilePath: \\Git\\MyScript\\用例\\WordToExcel.py
@Copyright (c) 2022 by ishuangjin, All Rights Reserved.
@Description: 
'''
import docx
from win32com import client
import pandas as pd
import os
"""
python -m pip install pywin32
pip install python-docx

ImportError: DLL load failed: 找不到指定的模块。
原因是pywin32的dll需要放到C:\Windows\System32目录下，
我的python默认目录为：D:\python\Lib\site-packages\pywin32_system32
该文件夹下有两个dll，“pythoncom35.dll” 和“pywintypes35.dll”，复制到C:\Windows\System32目录。
"""


def doc_docx(doc_case):
    '''
    @description: doc文件转化为docx文件
    @param  doc_case: Excel文档绝对路径
    @return 
    '''
    # 路径
    new_path = doc_case.replace('doc', 'docx')
    # 获取Word应用程序对象
    word_app = client.Dispatch('Word.Application')
    # 打开对应的Word文档
    doc = word_app.Documents.Open(doc_case)
    # 另存到docx_path中，12表示docx文件格式
    doc.SaveAs(new_path, 12)
    # 关闭Word文档
    doc.Close()
    # 退出软件
    word_app.Quit()
    return new_path


def save_data(doc_case, excel_case):
    try:
        doc = docx.Document(doc_case)
    except Exception:
        doc = docx.Document(doc_docx(doc_case))
    tables = doc.tables
    excel_path = excel_case
    list_ = []  # 初始化一个空列表，用来装后面的dict_
    for table in tables:  # 循环所有的表格列表
        dict_ = {}
        for i in range(0, len(table.rows)):  # 后续的内容格式不固定，所以循环获取
            try:
                dict_[table.cell(i, 0).text] = table.cell(i, 1).text
            except KeyError:
                print("如果代码64行报错，查看是否存在用例以外的表，删除掉即可")
            finally:
                list_.append(dict_)
    print("用例共{}条".format(len(tables)))
    df = pd.DataFrame(list_)
    # 重新设置一下行列索引的名字
    df.columns = ["用例名称", "测试说明", "前置条件", "测试步骤", "预期结果", "实际结果", "备注", "用例标识"]  # , "功能需求追踪", "用例编制人", "测试执行人", "测试结论"
    df = df.reindex(columns=['用例标识', '用例名称', '功能需求追踪', '测试说明', '前置条件', '测试步骤', '预期结果', '用例编制人', '实际结果', '测试结论', '测试执行人', '备注'])
    df.to_excel(excel_path, index=None)


def run(doc_file_name):
    # 路径操作
    script_path = os.path.dirname(__file__)  # 当前脚本的绝对路径
    doc_case = os.path.join(script_path, "WordCase", doc_file_name)
    excel_case = os.path.join(script_path, "ExcelCase", doc_file_name.replace(".docx", ".xlsx"))

    try:
        save_data(doc_case=doc_case, excel_case=excel_case)
        print("转化后的文件路径为:", excel_case)
    except PermissionError:
        print("Error:文件被占用,请关闭已打开的xlsx文件")
    else:
        pass


def main():
    # 填写Excel文件名
    doc_file_name = "TSF1.29测试用例.docx"  # 要操作的文件
    run(doc_file_name)


if __name__ == '__main__':
    main()
