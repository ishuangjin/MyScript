# 没有的先pip install 包名称
from docx import Document
from openpyxl import Workbook

# 指定要操作的doc文档
document = Document('腾讯云微服务平台TSF V1.29.4 测试用例.docx')

count = 0
tables = []
wb = Workbook()  # 创建Excel对象
ws = wb.active  # 创建工作簿

# 设置列数，可以指定列名称，有几列就设置几个，
# A对应列1，B对应列2，以此类推
# 只能处理列数一致的表格，不一致的请在word文档（转下行）
# 中处理好后，再运行程序
ws['A1'] = '国家'
ws['B1'] = '专利号'
ws['C1'] = '自己设置'
ws['D1'] = '懂？'
ws['E1'] = '5'
ws['F1'] = '6'
ws['G1'] = '7'
ws['H1'] = '8'
ws['I1'] = '9'
ws['J1'] = '10'
ws['K1'] = '11'
ws['L1'] = '12'

total = len(document.tables)
print("总共", total, "个表格等待处理，请喝杯咖啡等待许久...")
for index in range(0, total):
    table = []
    for row in document.tables[index].rows:
        line = []
        for grid in row.cells:
            line.append(grid.text)
        table.append(line)
        ws.append(line)
    count = count + 1
    print("第", count, "个表格正在处理...剩余", total - count + 1, "个表格", "\n")
    tables.append(table)
# 测试专用，测试前30条数据请打开注释
# if count == 30:
# break

wb.save("要保存xlsx的路径.xlsx")  # 保存到文件
print(tables)
print("表格处理完成...")
