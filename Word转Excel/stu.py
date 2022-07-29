from docx import Document
from openpyxl import Workbook
import pandas as pd

document = Document('腾讯云微服务平台TSF V1.29.4 测试用例.docx')
count = 0
tables = []
wb = Workbook()  # 创建Excel对象
ws = wb.active  # 创建工作簿
# 设置列数，可以指定列名称，有几列就设置几个，A对应列1，B对应列2，以此类推
ws['A1'] = '用例名称'
ws['B1'] = '前置条件'
ws['C1'] = '用例步骤'
ws['D1'] = '前置结果'
ws['E1'] = '备注'

total = len(document.tables)  # 返回列表，word文件中所有的表格数
print("总共", total, "个表格等待处理，请喝杯咖啡等待许久...")

for index in range(0, total):  # 循环word中的每个表格
    table = []  # 创建空列表用来储存表格
    for row in document.tables[index].rows:  # 第index个表格中遍历行
        line = []  # 创建空列表用来储存行
        for grid in row.cells:  # 遍历行的单元格
            line.append(grid.text)  # 将单元格内容添加到列表line中
        table.append(line)  # 给table追加列表line，line中已储存第row行的所有单元格信息
        ws.append(line)  # 给excel追加行
    count = count + 1
    print("第", count, "个表格正在处理...剩余", total - count + 1, "个表格", "\n")
    tables.append(table)  # 追加该word中的表格table到列表tables中，处理下一个表格

# wb.save("要保存xlsx的路径.xlsx")  # 保存到文件
# df = pd.read_excel("要保存xlsx的路径.xlsx")
# df_T = df.T  # 获得矩阵的转置
# df_T.to_excel('要保存xlsx的路径_1.xlsx', sheet_name='测试用例转换')  # 保存文件
# print("表格处理完成...")
